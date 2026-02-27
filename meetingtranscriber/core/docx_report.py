# meetingtranscriber/core/docx_report.py
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from meetingtranscriber.core.summarizer import SmartSummary


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(it)


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_docx_report(
    out_path: Path,
    created_at: str,
    source_path: str,
    duration_sec: float,
    full_text: str,
    summary: SmartSummary,
    app_name: str = "ALSA",
) -> Path:
    doc = Document()

    _add_title(doc, f"{app_name} — تقرير تفريغ وتلخيص اجتماع")

    _add_paragraph(doc, f"التاريخ: {created_at}")
    _add_paragraph(doc, f"الملف: {source_path}")
    if duration_sec:
        _add_paragraph(doc, f"المدة التقريبية: {duration_sec:.1f} ثانية")

    doc.add_paragraph("")

    _add_h2(doc, "ملخص سريع (نقاط)")
    if summary.bullets:
        _add_bullets(doc, summary.bullets)
    else:
        _add_paragraph(doc, "لا يوجد نقاط كافية.")

    doc.add_paragraph("")

    _add_h2(doc, "المحاور (حسب الموضوع)")
    if summary.topics:
        for topic, items in summary.topics.items():
            _add_h2(doc, f"- {topic}")
            _add_bullets(doc, items)
            doc.add_paragraph("")
    else:
        _add_paragraph(doc, "لا يوجد محاور مستخرجة.")

    doc.add_paragraph("")

    _add_h2(doc, "حسب المتحدثين (إن توفر)")
    if summary.speakers:
        for sp, items in summary.speakers.items():
            _add_h2(doc, f"- {sp}")
            _add_bullets(doc, items)
            doc.add_paragraph("")
    else:
        _add_paragraph(doc, "لم يتم التعرف على متحدثين (لا توجد Labels في النص).")

    doc.add_paragraph("")

    _add_h2(doc, "القرارات")
    if summary.decisions:
        _add_bullets(doc, summary.decisions)
    else:
        _add_paragraph(doc, "لا توجد قرارات واضحة في النص.")

    doc.add_paragraph("")

    _add_h2(doc, "المهام (To-Do)")
    if summary.tasks:
        _add_bullets(doc, summary.tasks)
    else:
        _add_paragraph(doc, "لا توجد مهام واضحة في النص.")

    doc.add_paragraph("")

    _add_h2(doc, "أرقام/تواريخ/مؤشرات مذكورة")
    if summary.numbers_dates:
        _add_bullets(doc, summary.numbers_dates)
    else:
        _add_paragraph(doc, "لا يوجد.")

    doc.add_paragraph("")
    _add_h2(doc, "النص الكامل")
    _add_paragraph(doc, full_text or "")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path