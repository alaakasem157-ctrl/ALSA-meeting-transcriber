from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt


def export_summary_docx(
    out_path: Path,
    title: str,
    overview: str,
    key_points: List[str],
    decisions: List[str],
    tasks: List[Dict[str, str]],
    speakers: List[Dict[str, Any]],
    source_path: str = "",
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_default_font(doc, "Segoe UI")

    doc.add_heading(title, level=1)

    meta = []
    if source_path:
        meta.append(f"Source: {source_path}")
    meta.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(" | ".join(meta))

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(overview or "")

    doc.add_heading("Key Points", level=2)
    if key_points:
        for p in key_points:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("-")

    doc.add_heading("Decisions", level=2)
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("-")

    doc.add_heading("Tasks", level=2)
    if tasks:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Task"
        hdr[1].text = "Owner"
        hdr[2].text = "Due"

        for t in tasks:
            row = table.add_row().cells
            row[0].text = str(t.get("task") or "")
            row[1].text = str(t.get("owner") or "")
            row[2].text = str(t.get("due") or "")
    else:
        doc.add_paragraph("-")

    if speakers:
        doc.add_heading("By Speaker", level=2)
        for sp in speakers:
            name = str(sp.get("name") or "").strip()
            notes = sp.get("notes") or []
            if not name:
                continue
            doc.add_heading(name, level=3)
            if notes:
                for n in notes:
                    doc.add_paragraph(str(n), style="List Bullet")
            else:
                doc.add_paragraph("-")

    doc.save(str(out_path))
    return out_path


def _set_default_font(doc: Document, font_name: str) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(11)