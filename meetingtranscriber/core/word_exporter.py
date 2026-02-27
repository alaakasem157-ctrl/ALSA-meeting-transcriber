from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from meetingtranscriber.core.paths import outputs_dir


FONT_NAME = "WinSoft Pro"
BASE_FONT_SIZE = 14


def _set_paragraph_rtl(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _set_run_font(run, size_pt: int = BASE_FONT_SIZE, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bool(bold)
    rPr = run._r.get_or_add_rPr()

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.append(rFonts)

    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def _add_header_logo(doc: Document, logo_path: Optional[str]) -> None:
    if not logo_path:
        return
    lp = Path(str(logo_path))
    if not lp.exists():
        return

    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        run = p.add_run()
        run.add_picture(str(lp), width=Inches(2.2))
    except Exception:
        pass
    _set_paragraph_rtl(p)


def _add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rtl(p)
    r = p.add_run(title)
    _set_run_font(r, size_pt=18, bold=True)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    r = p.add_run(text)
    _set_run_font(r, size_pt=16, bold=True)


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    r = p.add_run(text or "")
    _set_run_font(r, size_pt=BASE_FONT_SIZE, bold=False)


def _add_bullets(doc: Document, items: List[str]) -> None:
    for it in items or []:
        t = str(it).strip()
        if not t:
            continue
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_rtl(p)
        r = p.add_run(t)
        _set_run_font(r, size_pt=BASE_FONT_SIZE, bold=False)


def _normalize_call(args: Tuple[Any, ...], kwargs: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(kwargs)

    out_path = out.pop("out_path", None) or out.pop("output_path", None)
    header_logo_path = out.pop("header_logo_path", None) or out.pop("header_image_path", None)
    transcript_text = out.pop("transcript_text", None) or out.pop("transcript", None) or out.pop("text", None)
    summary_text = out.pop("summary_text", None) or out.pop("summary", None)

    topics = out.pop("topics", None)
    decisions = out.pop("decisions", None)
    tasks = out.pop("tasks", None) or out.pop("action_items", None)

    source_path = out.pop("source_path", None) or out.pop("source", None)
    title = out.pop("title", None)

    meta = out.pop("meta", None) or {}

    a = list(args)

    if a and out_path is None:
        s0 = str(a[0])
        if s0.lower().endswith(".docx"):
            out_path = s0
            a = a[1:]

    if a and title is None:
        title = str(a[0])
        a = a[1:]

    if a and transcript_text is None:
        transcript_text = str(a[0])
        a = a[1:]

    if a and summary_text is None:
        summary_text = str(a[0])
        a = a[1:]

    if title is None:
        title = "محضر اجتماع"

    if transcript_text is None:
        transcript_text = ""

    if summary_text is None:
        summary_text = ""

    def _as_list(v) -> List[str]:
        if v is None:
            return []
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        if isinstance(v, str):
            lines = [ln.strip(" -•\t") for ln in v.splitlines()]
            return [ln for ln in lines if ln]
        return [str(v).strip()] if str(v).strip() else []

    return {
        "out_path": out_path,
        "title": str(title),
        "transcript_text": str(transcript_text),
        "summary_text": str(summary_text),
        "topics": _as_list(topics),
        "decisions": _as_list(decisions),
        "tasks": _as_list(tasks),
        "source_path": str(source_path or ""),
        "meta": dict(meta) if isinstance(meta, dict) else {},
        "header_logo_path": str(header_logo_path) if header_logo_path else None,
    }


def export_meeting_docx(*args: Any, **kwargs: Any) -> str:
    data = _normalize_call(args, kwargs)

    out_path = data["out_path"]
    if not out_path:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = str(outputs_dir() / f"meeting_{stamp}.docx")

    outp = Path(out_path).resolve()
    outp.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BASE_FONT_SIZE)

    _add_header_logo(doc, data["header_logo_path"])

    _add_title(doc, data["title"])

    meta_lines = []
    meta_lines.append(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if data["source_path"]:
        meta_lines.append(f"المصدر: {Path(data['source_path']).name}")

    for k, v in (data["meta"] or {}).items():
        if v is None:
            continue
        vv = str(v).strip()
        if vv:
            meta_lines.append(f"{k}: {vv}")

    if meta_lines:
        p = doc.add_paragraph()
        _set_paragraph_rtl(p)
        for i, line in enumerate(meta_lines):
            r = p.add_run(line)
            _set_run_font(r, size_pt=11, bold=False)
            if i != len(meta_lines) - 1:
                r.add_break()

    doc.add_paragraph()

    _add_heading(doc, "الملخص")
    _add_paragraph(doc, data["summary_text"].strip() if data["summary_text"] else "")

    doc.add_paragraph()

    _add_heading(doc, "محاور الاجتماع")
    if data["topics"]:
        _add_bullets(doc, data["topics"])
    else:
        _add_paragraph(doc, "غير محدد")

    doc.add_paragraph()

    _add_heading(doc, "القرارات")
    if data["decisions"]:
        _add_bullets(doc, data["decisions"])
    else:
        _add_paragraph(doc, "لا يوجد قرارات واضحة ضمن النص.")

    doc.add_paragraph()

    _add_heading(doc, "المهام")
    if data["tasks"]:
        _add_bullets(doc, data["tasks"])
    else:
        _add_paragraph(doc, "لا يوجد مهام واضحة ضمن النص.")

    doc.add_paragraph()

    _add_heading(doc, "النص المفرغ")
    _add_paragraph(doc, data["transcript_text"].strip() if data["transcript_text"] else "")

    doc.save(str(outp))
    return str(outp)