from pathlib import Path
from datetime import datetime

from docx import Document


def export_meeting_docx(
    output_path: Path,
    report_title: str,
    source_name: str,
    full_text: str,
    summary_data: dict,
):
    """
    Export a clean DOCX report.
    (Simple layout, ready for SVU template later.)
    """
    doc = Document()

    # Title
    doc.add_heading(report_title, level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if source_name:
        doc.add_paragraph(f"Source: {source_name}")

    doc.add_paragraph("")

    # Summary
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(summary_data.get("summary", ""))

    # Key Points
    doc.add_heading("Key Points", level=2)
    for p in summary_data.get("key_points", []):
        doc.add_paragraph(p, style="List Bullet")

    # Action Items
    doc.add_heading("Action Items", level=2)
    actions = summary_data.get("action_items", [])
    if actions:
        for a in actions:
            doc.add_paragraph(a, style="List Number")
    else:
        doc.add_paragraph("—")

    # Transcript
    doc.add_heading("Full Transcript", level=2)
    doc.add_paragraph(full_text or "")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))